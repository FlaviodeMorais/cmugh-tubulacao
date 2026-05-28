import base64
import hashlib
import io
import json
import re
import textwrap
from dataclasses import dataclass, asdict
from datetime import date
from typing import List

import pandas as pd
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from PIL import Image
import requests

import streamlit as st

# Credenciais do admin master (senha armazenada como SHA-256)
_ADMIN_USER = "Admin"
_ADMIN_HASH = hashlib.sha256("979101Fm$".encode()).hexdigest()


def _verificar_admin(usuario: str, senha: str) -> bool:
    return (usuario == _ADMIN_USER and
            hashlib.sha256(senha.encode()).hexdigest() == _ADMIN_HASH)

DISCIPLINAS = [
    "Tubulação", "Dinâmicos", "Estáticos", "Civil", "Estruturas Metálicas",
    "Elétrica", "Instrumentação", "Telecom", "Automação", "HVAC",
    "Comissionamento", "Qualidade", "SMS", "Contratual",
]

# ─────────────────────────── DATACLASS ────────────────────────────

@dataclass
class RegistroCM:
    id: int
    tenant: str
    data_registro: str
    obra: str
    frente_servico: str
    disciplina: str
    atividade: str
    equipe: str
    responsavel: str
    fiscal: str
    status: str
    impacto_rdo: str
    observacoes: str
    evidencias: str = ""
    chave: str = ""

# ─────────────────────────── BANCO DE DADOS (Supabase REST) ──────────────────

def _sb_url() -> str:
    return st.secrets.get("supabase", {}).get("url", "") + "/rest/v1"

def _sb_key() -> str:
    return st.secrets.get("supabase", {}).get("key", "")

def _headers(prefer_rep: bool = False) -> dict:
    key = _sb_key()
    h = {"apikey": key, "Authorization": f"Bearer {key}", "Content-Type": "application/json"}
    if prefer_rep:
        h["Prefer"] = "return=representation"
    return h

class _SupabaseError(Exception):
    pass


def _handle_request_error(e: Exception) -> None:
    if isinstance(e, requests.exceptions.Timeout):
        msg = ("⏱️ **Tempo limite excedido ao conectar com o banco de dados.**\n\n"
               "O projeto Supabase pode estar **pausado** (plano gratuito pausa após 7 dias sem uso).\n\n"
               "👉 Acesse [app.supabase.com](https://app.supabase.com), abra o projeto e clique em **Restore project**.")
    elif isinstance(e, requests.exceptions.ConnectionError):
        msg = "🔌 **Sem conexão com o banco de dados.** Verifique sua internet e tente novamente."
    else:
        msg = f"❌ **Erro ao comunicar com o banco de dados:** {e}"
    raise _SupabaseError(msg)


def _get(table: str, params: dict | None = None) -> list:
    try:
        r = requests.get(f"{_sb_url()}/{table}", headers=_headers(), params=params, timeout=15)
        r.raise_for_status()
        return r.json()
    except (requests.exceptions.Timeout, requests.exceptions.ConnectionError) as e:
        _handle_request_error(e)
    except requests.exceptions.HTTPError as e:
        raise _SupabaseError(f"❌ Erro HTTP {e.response.status_code}: {e.response.text[:200]}") from e

def _post(table: str, data: dict) -> list:
    try:
        r = requests.post(f"{_sb_url()}/{table}", headers=_headers(prefer_rep=True), json=data, timeout=15)
        r.raise_for_status()
        return r.json()
    except (requests.exceptions.Timeout, requests.exceptions.ConnectionError) as e:
        _handle_request_error(e)
    except requests.exceptions.HTTPError as e:
        raise _SupabaseError(f"❌ Erro HTTP {e.response.status_code}: {e.response.text[:200]}") from e

def _patch(table: str, params: dict, data: dict) -> None:
    try:
        r = requests.patch(f"{_sb_url()}/{table}", headers=_headers(), params=params, json=data, timeout=15)
        r.raise_for_status()
    except (requests.exceptions.Timeout, requests.exceptions.ConnectionError) as e:
        _handle_request_error(e)
    except requests.exceptions.HTTPError as e:
        raise _SupabaseError(f"❌ Erro HTTP {e.response.status_code}: {e.response.text[:200]}") from e

def _delete(table: str, params: dict) -> None:
    try:
        r = requests.delete(f"{_sb_url()}/{table}", headers=_headers(), params=params, timeout=15)
        r.raise_for_status()
    except (requests.exceptions.Timeout, requests.exceptions.ConnectionError) as e:
        _handle_request_error(e)
    except requests.exceptions.HTTPError as e:
        raise _SupabaseError(f"❌ Erro HTTP {e.response.status_code}: {e.response.text[:200]}") from e


def init_db() -> None:
    pass  # tabelas já criadas no Supabase


def _eq(value) -> str:
    return f"eq.{value}"


# ── contratos ──

@st.cache_data(ttl=300, show_spinner=False)
def listar_contratos() -> List[str]:
    rows = _get("contratos", {"select": "numero", "order": "numero.asc"})
    return [r["numero"] for r in rows]


def criar_contrato(numero: str, senha: str, identificador: str = "") -> bool:
    try:
        _post("contratos", {"numero": numero, "senha_admin": senha, "identificador": identificador})
        listar_contratos.clear()
        return True
    except Exception:
        return False


@st.cache_data(ttl=300, show_spinner=False)
def obter_identificador(numero: str) -> str:
    rows = _get("contratos", {"select": "identificador", "numero": _eq(numero)})
    return rows[0].get("identificador", "") if rows else ""


def atualizar_identificador(numero: str, identificador: str) -> None:
    _patch("contratos", {"numero": _eq(numero)}, {"identificador": identificador})
    obter_identificador.clear()


def verificar_senha(numero: str, senha: str) -> bool:
    rows = _get("contratos", {"select": "id", "numero": _eq(numero), "senha_admin": _eq(senha)})
    return len(rows) > 0


def excluir_contrato(numero: str) -> None:
    _delete("contratos", {"numero": _eq(numero)})
    listar_contratos.clear()


# ── fiscais ──

@st.cache_data(ttl=120, show_spinner=False)
def listar_fiscais(contrato: str) -> list:
    return _get("fiscais", {"select": "*", "contrato": _eq(contrato), "order": "nome.asc"})


def adicionar_fiscal(contrato, nome, chave, disciplina, email="") -> None:
    _post("fiscais", {"contrato": contrato, "nome": nome, "chave": chave, "disciplina": disciplina, "email": email})
    listar_fiscais.clear()


def excluir_fiscal(fid: int) -> None:
    _delete("fiscais", {"id": _eq(fid)})
    listar_fiscais.clear()


# ── unidades ──

@st.cache_data(ttl=120, show_spinner=False)
def listar_unidades(contrato: str) -> list:
    return _get("unidades", {"select": "*", "contrato": _eq(contrato), "order": "nome.asc"})


def adicionar_unidade(contrato: str, nome: str) -> None:
    _post("unidades", {"contrato": contrato, "nome": nome})
    listar_unidades.clear()


def excluir_unidade(uid: int) -> None:
    _delete("unidades", {"id": _eq(uid)})
    listar_unidades.clear()


# ── registros ──

@st.cache_data(ttl=60, show_spinner=False)
def carregar_lista(tenant: str) -> List[RegistroCM]:
    rows = _get("registros_cm", {"select": "*", "tenant": _eq(tenant), "order": "id.desc"})
    return [RegistroCM(**r) for r in rows]


def salvar_registro(registro: RegistroCM) -> None:
    _post("registros_cm", {
        "tenant":         registro.tenant,
        "data_registro":  registro.data_registro,
        "obra":           registro.obra,
        "frente_servico": registro.frente_servico,
        "disciplina":     registro.disciplina,
        "atividade":      registro.atividade,
        "equipe":         registro.equipe,
        "responsavel":    registro.responsavel,
        "fiscal":         registro.fiscal,
        "status":         registro.status,
        "impacto_rdo":    registro.impacto_rdo,
        "observacoes":    registro.observacoes,
        "evidencias":     registro.evidencias,
        "chave":          registro.chave,
    })
    carregar_lista.clear()


def excluir_registro(rid: int) -> None:
    _delete("registros_cm", {"id": _eq(rid)})
    carregar_lista.clear()
    _pdf_registro_cached.clear()


def excluir_registros(ids: list[int]) -> None:
    for rid in ids:
        _delete("registros_cm", {"id": _eq(rid)})
    carregar_lista.clear()
    _pdf_registro_cached.clear()

# ─────────────────────────── HELPERS VISUAIS ──────────────────────

THUMB_SIZE = 200


def img_thumb(data: bytes) -> bytes:
    img = Image.open(io.BytesIO(data)).convert("RGB")
    w, h = img.size
    m = min(w, h)
    img = img.crop(((w - m) // 2, (h - m) // 2, (w + m) // 2, (h + m) // 2))
    img = img.resize((THUMB_SIZE, THUMB_SIZE), Image.LANCZOS)
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=85)
    return buf.getvalue()


def fmt_data(iso: str) -> str:
    try:
        a, m, d = iso.split("-")
        return f"{d}/{m}/{a}"
    except Exception:
        return iso


def id_registro(r) -> str:
    data = fmt_data(r.data_registro).replace("/", "")
    disc = (r.disciplina or "").replace(" ", "")
    fiscal = (r.fiscal or "").strip()
    partes = [p for p in ["RO", f"{r.id:04d}", r.obra, disc, data, fiscal] if p]
    return "-".join(partes)


def _parse_evidencias(evidencias_json: str) -> list:
    try:
        items = json.loads(evidencias_json)
    except Exception:
        return []
    result = []
    for item in items:
        if isinstance(item, dict):
            result.append({"foto": base64.b64decode(item["foto"]), "legenda": item.get("legenda", "")})
        else:
            result.append({"foto": base64.b64decode(item), "legenda": ""})
    return result


def _img_centralizada(data: bytes, legenda: str = "") -> None:
    b64 = base64.b64encode(data).decode()
    cap = (
        f'<p style="text-align:center;font-size:0.78em;color:#aaa;margin:4px 0 8px">{legenda}</p>'
        if legenda else ""
    )
    st.markdown(
        f'<div style="display:flex;flex-direction:column;align-items:center">'
        f'<img src="data:image/jpeg;base64,{b64}" width="{THUMB_SIZE}">'
        f'{cap}</div>',
        unsafe_allow_html=True,
    )


def exibir_grid_evidencias(evidencias_json: str) -> None:
    if not evidencias_json:
        return
    items = _parse_evidencias(evidencias_json)
    if not items:
        return
    thumbs = [{"t": img_thumb(i["foto"]), "l": i["legenda"]} for i in items]
    if len(thumbs) == 1:
        _img_centralizada(thumbs[0]["t"], thumbs[0]["l"])
    else:
        for linha in range(0, len(thumbs), 2):
            cols = st.columns(2)
            for col_idx, img_idx in enumerate(range(linha, min(linha + 2, len(thumbs)))):
                with cols[col_idx]:
                    _img_centralizada(thumbs[img_idx]["t"], thumbs[img_idx]["l"])


def para_exibicao(registros: List[RegistroCM]) -> list:
    rows = []
    for r in registros:
        d = asdict(r)
        d["data_registro"] = fmt_data(d["data_registro"])
        try:
            n = len(json.loads(d.get("evidencias") or "[]"))
        except Exception:
            n = 0
        d["evidencias"] = f"{n} foto(s)" if n else ""
        rows.append(d)
    return rows


def exibir_cards(registros: List[RegistroCM], contrato: str = "", empreendimento: str = "") -> None:
    for r in registros:
        rid = id_registro(r)
        with st.expander(rid, expanded=False):
            st.checkbox("Selecionar para PDF", key=f"sel_{r.id}")
            col1, col2 = st.columns(2)
            with col1:
                st.markdown(f"**Frente:** {r.frente_servico}")
                st.markdown(f"**Disciplina:** {r.disciplina}")
                st.markdown(f"**Atividade:** {r.atividade}")
            with col2:
                st.markdown(f"**Equipe:** {r.equipe or '—'}")
                st.markdown(f"**Responsável:** {r.responsavel or '—'}")
                st.markdown(f"**Status:** {r.status}")
                st.markdown(f"**Classificação:** {r.impacto_rdo}")
            if r.observacoes:
                st.markdown(f"**Obs:** {r.observacoes}")
            exibir_grid_evidencias(r.evidencias)
            if contrato:
                _share_btn(
                    "Compartilhar PDF",
                    _pdf_registro_cached(
                        r.id, r.evidencias, contrato, empreendimento,
                        r.tenant, r.data_registro, r.obra, r.frente_servico,
                        r.disciplina, r.atividade, r.equipe, r.responsavel,
                        r.fiscal, r.status, r.impacto_rdo, r.observacoes, r.chave,
                    ),
                    f"{rid}.pdf",
                    "application/pdf",
                    f"pdf_ro_{r.id}",
                )

# ─────────────────────────── EXPORTAÇÕES ──────────────────────────

def _to_excel(registros) -> bytes:
    df = pd.DataFrame(para_exibicao(registros))
    df = df.drop(columns=["id", "tenant", "evidencias"], errors="ignore")
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="RDOe")
    return buf.getvalue()


def _cabecalho_word(doc: Document, contrato: str, empreendimento: str = "") -> None:
    t = doc.add_paragraph()
    run = t.add_run("RO - Registro de Ocorrências")
    run.bold = True
    run.font.size = Pt(14)
    sub = doc.add_paragraph()
    _ide = empreendimento or "SRGE/SI-III/HDTON/CMUGH"
    sub.add_run(f"{_ide}  |  Contrato: {contrato}").font.size = Pt(10)
    doc.add_paragraph()


def _to_word(registros, contrato: str, empreendimento: str = "") -> bytes:
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin = Cm(2)
        sec.right_margin = Cm(2)
    _cabecalho_word(doc, contrato, empreendimento)
    for r in registros:
        doc.add_heading(id_registro(r), level=2)
        tabela = doc.add_table(rows=5, cols=2)
        tabela.style = "Table Grid"
        dados = [
            ("Fiscal de Campo", r.fiscal or "—", "", ""),
            ("Frente de Serviço", r.frente_servico, "Equipe", r.equipe or "—"),
            ("Responsável Contratada", r.responsavel or "—", "", ""),
            ("Disciplina", r.disciplina, "Status", r.status),
            ("Atividade Executada", r.atividade, "Classificação", r.impacto_rdo),
        ]
        for i, (l1, v1, l2, v2) in enumerate(dados):
            tabela.cell(i, 0).text = f"{l1}: {v1}"
            tabela.cell(i, 1).text = f"{l2}: {v2}"
        if r.observacoes:
            doc.add_paragraph(f"Observações: {r.observacoes}")
        items = _parse_evidencias(r.evidencias)
        if items:
            doc.add_paragraph("Evidências:").runs[0].bold = True
            thumbs = [(img_thumb(i["foto"]), i["legenda"]) for i in items]
            for par in range(0, len(thumbs), 2):
                tbl = doc.add_table(rows=2, cols=min(2, len(thumbs) - par))
                for col_idx in range(len(tbl.columns)):
                    thumb_data, legenda = thumbs[par + col_idx]
                    cell_p = tbl.cell(0, col_idx).paragraphs[0]
                    cell_p.add_run().add_picture(io.BytesIO(thumb_data), width=Cm(6))
                    leg_p = tbl.cell(1, col_idx).paragraphs[0]
                    leg_run = leg_p.add_run(legenda)
                    leg_run.font.size = Pt(8)
                    leg_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        doc.add_paragraph()
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class _PDF(FPDF):
    def __init__(self, contrato: str, empreendimento: str = ""):
        super().__init__()
        self._contrato = contrato
        self._empreendimento = empreendimento or "SRGE/SI-III/HDTON/CMUGH"

    def header(self):
        self.set_font("Helvetica", "B", 12)
        self.cell(0, 8, "RO - Registro de Ocorrencias", align="L", new_x="LMARGIN", new_y="NEXT")
        self.set_font("Helvetica", "", 9)
        self.cell(0, 6, f"{self._empreendimento}  |  Contrato: {self._contrato}", align="L", new_x="LMARGIN", new_y="NEXT")
        self.ln(2)

    def footer(self):
        self.set_y(-12)
        self.set_font("Helvetica", "I", 8)
        self.cell(0, 8, f"Pag. {self.page_no()}", align="C")


def _to_pdf(registros, contrato: str, empreendimento: str = "") -> bytes:
    pdf = _PDF(contrato, empreendimento)
    pdf.set_auto_page_break(auto=True, margin=15)
    for r in registros:
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 7, id_registro(r), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)
        for label, valor in [
            ("Fiscal", r.fiscal or "-"),
            ("Frente", r.frente_servico), ("Disciplina", r.disciplina),
            ("Atividade", r.atividade), ("Equipe", r.equipe or "-"),
            ("Responsável", r.responsavel or "-"),
            ("Status", r.status), ("Classificação", r.impacto_rdo),
        ]:
            pdf.set_font("Helvetica", "B", 9)
            pdf.write(6, f"{label}: ")
            pdf.set_font("Helvetica", "", 9)
            pdf.write(6, valor)
            pdf.ln(6)
        if r.observacoes:
            pdf.set_font("Helvetica", "B", 9)
            pdf.write(6, "Obs: ")
            pdf.set_font("Helvetica", "", 9)
            pdf.write(6, r.observacoes)
            pdf.ln(6)
        items = _parse_evidencias(r.evidencias)
        if items:
            thumbs = [(img_thumb(i["foto"]), i["legenda"]) for i in items]
            col_w = 80
            row_h = col_w + 14  # 80mm imagem + 14mm legenda
            page_bottom = pdf.h - pdf.b_margin
            num_rows = (len(thumbs) + 1) // 2
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 9)
            pdf.cell(0, 6, "Evidencias:", new_x="LMARGIN", new_y="NEXT")
            x0 = pdf.get_x()
            for row_idx in range(num_rows):
                # só quebra página se a imagem em si não couber (legenda pode sobrar um pouco)
                if pdf.get_y() + col_w > page_bottom:
                    pdf.add_page()
                y_row = pdf.get_y()
                for col_idx in range(2):
                    thumb_idx = row_idx * 2 + col_idx
                    if thumb_idx >= len(thumbs):
                        break
                    thumb_data, legenda = thumbs[thumb_idx]
                    x = x0 + col_idx * (col_w + 5)
                    pdf.image(io.BytesIO(thumb_data), x=x, y=y_row, w=col_w)
                    pdf.set_xy(x, y_row + col_w + 1)
                    pdf.set_font("Helvetica", "I", 7)
                    lines = textwrap.wrap(legenda, width=42)[:2]
                    pdf.multi_cell(col_w, 3.5, "\n".join(lines), align="C")
                pdf.set_xy(x0, y_row + row_h)
    return bytes(pdf.output())

@st.cache_data(show_spinner=False)
def _pdf_registro_cached(r_id: int, evidencias: str, contrato: str, empreendimento: str,
                         tenant: str, data_registro: str, obra: str, frente_servico: str,
                         disciplina: str, atividade: str, equipe: str, responsavel: str,
                         fiscal: str, status: str, impacto_rdo: str, observacoes: str,
                         chave: str) -> bytes:
    r = RegistroCM(id=r_id, tenant=tenant, data_registro=data_registro, obra=obra,
                   frente_servico=frente_servico, disciplina=disciplina, atividade=atividade,
                   equipe=equipe, responsavel=responsavel, fiscal=fiscal, status=status,
                   impacto_rdo=impacto_rdo, observacoes=observacoes,
                   evidencias=evidencias, chave=chave)
    return _to_pdf([r], contrato, empreendimento)

# ─────────────────────────── COMPARTILHAMENTO ─────────────────────

def _share_btn(label: str, data: bytes, filename: str, mime: str, key: str) -> None:
    """Botão que usa Web Share API no mobile; fallback download no desktop."""
    b64 = base64.b64encode(data).decode()
    fid = re.sub(r"\W", "_", key)
    html = f"""
<style>
#sbtn_{fid}{{background:#007A33;color:#fff;border:none;padding:8px 16px;
border-radius:6px;cursor:pointer;font-size:14px;width:100%;
font-weight:700;font-family:sans-serif;letter-spacing:0.04em;text-transform:uppercase;}}
#sbtn_{fid}:hover{{background:#005C26;}}
</style>
<button id="sbtn_{fid}" onclick="sf_{fid}()">{label}</button>
<script>
async function sf_{fid}(){{
  const raw=atob('{b64}');
  const buf=new Uint8Array(raw.length);
  for(let i=0;i<raw.length;i++) buf[i]=raw.charCodeAt(i);
  const blob=new Blob([buf],{{type:'{mime}'}});
  const file=new File([blob],'{filename}',{{type:'{mime}'}});
  if(navigator.canShare&&navigator.canShare({{files:[file]}})){{
    try{{await navigator.share({{files:[file],title:'{filename}'}});return;}}
    catch(e){{if(e.name==='AbortError')return;}}
  }}
  const url=URL.createObjectURL(blob);
  const w=window.open(url,'_blank');
  if(!w){{
    const a=document.createElement('a');
    a.href=url;a.download='{filename}';
    document.body.appendChild(a);a.click();document.body.removeChild(a);
  }}
  setTimeout(()=>URL.revokeObjectURL(url),90000);
}}
</script>"""
    st.markdown(html, unsafe_allow_html=True)


# ─────────────────────────── INICIALIZAÇÃO ────────────────────────

init_db()
st.set_page_config(page_title="RO - Registro de Ocorrências", layout="centered")
st.markdown("""
<style>
/* ── ocultar chrome do Streamlit ── */
#MainMenu {visibility: hidden;}
header[data-testid="stHeader"] {visibility: hidden;}
footer {visibility: hidden !important; display: none !important;}
[data-testid="stToolbar"] {display: none !important;}
[data-testid="stBottom"] {display: none !important;}
[data-testid="manage-app-button"] {display: none !important;}
[class*="viewerBadge"] {display: none !important;}
[class*="manage"] {display: none !important;}

/* ── tema branco ── */
html, body, [data-testid="stApp"], [data-testid="stAppViewContainer"],
[data-testid="stMainBlockContainer"], .main, .block-container {
    background-color: #FFFFFF !important;
    color: #0D0D0D !important;
}
[data-testid="stSidebar"] {background-color: #F5F5F5 !important;}

/* texto geral */
p, span, label, div, li, h1, h2, h3, h4, h5, h6,
[data-testid="stMarkdownContainer"] * {
    color: #0D0D0D !important;
}

/* inputs */
input, textarea, select,
[data-baseweb="input"] input,
[data-baseweb="textarea"] textarea,
[data-baseweb="select"] div {
    background-color: #F8F8F8 !important;
    color: #0D0D0D !important;
    border-color: #CCCCCC !important;
}

/* selectbox / dropdown */
[data-baseweb="popover"] [role="option"],
[data-baseweb="menu"] li {
    background-color: #FFFFFF !important;
    color: #0D0D0D !important;
}

/* expander */
[data-testid="stExpander"] {
    background-color: #FAFAFA !important;
    border-color: #E0E0E0 !important;
}
[data-testid="stExpander"] summary p {
    color: #0D0D0D !important;
}

/* container com borda */
[data-testid="stVerticalBlock"] > div[data-testid="stVerticalBlockBorderWrapper"] {
    background-color: #FAFAFA !important;
    border-color: #E0E0E0 !important;
}

/* botões gerais */
button[kind="secondary"], button[data-testid="baseButton-secondary"] {
    background-color: #F0F0F0 !important;
    color: #0D0D0D !important;
    border-color: #CCCCCC !important;
}

/* botão primário — verde Petrobras */
button[data-testid="baseButton-primary"],
button[kind="primary"] {
    background-color: #007A33 !important;
    border-color: #007A33 !important;
    color: #FFFFFF !important;
    font-weight: 700 !important;
    letter-spacing: 0.04em !important;
    text-transform: uppercase !important;
}
button[data-testid="baseButton-primary"]:hover,
button[kind="primary"]:hover {
    background-color: #005C26 !important;
    border-color: #005C26 !important;
    color: #FFFFFF !important;
}
button[data-testid="baseButton-primary"] p,
button[kind="primary"] p {
    color: #FFFFFF !important;
}

/* botão de menu — só ícone, sem caixa */
.icon-btn button,
.icon-btn button:hover,
.icon-btn button:focus,
.icon-btn button:active {
    background: transparent !important;
    border: none !important;
    box-shadow: none !important;
    color: #0D0D0D !important;
    font-size: 1.3rem !important;
    padding: 0 !important;
    min-height: unset !important;
    line-height: 1 !important;
}
.icon-btn button:hover { opacity: 0.6 !important; }

/* Linha título + botão: impede empilhamento no mobile */
[data-testid="stMarkdownContainer"] + [data-testid="stHorizontalBlock"] {
    flex-wrap: nowrap !important;
    align-items: center !important;
}
[data-testid="stMarkdownContainer"] + [data-testid="stHorizontalBlock"]
  [data-testid="stColumn"]:last-child {
    flex: 0 0 auto !important;
    min-width: 50px !important;
}
</style>
""", unsafe_allow_html=True)

# ─────────────────────────── ESTADO ADMIN ────────────────────────

if "admin_logado" not in st.session_state:
    st.session_state.admin_logado = False
if "show_admin" not in st.session_state:
    st.session_state.show_admin = st.query_params.get("adm") == "1"

# Mostra Manage App apenas se admin estiver autenticado
if st.session_state.admin_logado:
    st.markdown("""
<style>
[data-testid="stBottom"] {display: flex !important;}
[data-testid="manage-app-button"] {display: flex !important;}
[class*="viewerBadge"] {display: block !important;}
</style>
""", unsafe_allow_html=True)

# ─────────────────────────── CABEÇALHO ───────────────────────────

@st.cache_resource
def _banner_b64() -> str:
    from pathlib import Path
    p = Path(__file__).parent / "header.jpg"
    if not p.exists():
        return ""
    img = Image.open(p).convert("RGB")
    w, h = img.size
    target_h = max(1, w // 6)
    if h > target_h:
        top = (h - target_h) // 2
        img = img.crop((0, top, w, top + target_h))
    img = img.resize((1200, 200), Image.LANCZOS)
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=82)
    return base64.b64encode(buf.getvalue()).decode()

@st.cache_resource
def _logo_b64() -> str:
    from pathlib import Path
    base = Path(__file__).parent
    for nome in ("PETROBRAS.jpg", "petrobras.jpg", "petrobras.png", "logo.png", "logo.jpg"):
        p = base / nome
        if p.exists():
            mime = "png" if nome.endswith(".png") else "jpeg"
            return mime, base64.b64encode(p.read_bytes()).decode()
    return "", ""

_b64 = _banner_b64()
_logo_mime, _logo_data = _logo_b64()

_banner_img = (
    f'<img src="data:image/jpeg;base64,{_b64}" '
    f'style="width:100%;height:160px;object-fit:cover;object-position:center 40%;display:block">'
    if _b64 else '<div style="height:160px;background:#007A33"></div>'
)
_logo_img = (
    f'<img src="data:image/{_logo_mime};base64,{_logo_data}" '
    f'style="height:34px;object-fit:contain;display:block">'
    if _logo_data else ""
)

if _logo_img:
    st.markdown(f'<div style="margin-bottom:16px">{_logo_img}</div>', unsafe_allow_html=True)

st.markdown(
    f'<div style="border-radius:8px;overflow:hidden;margin-bottom:8px">{_banner_img}</div>',
    unsafe_allow_html=True,
)

# Título (esquerda) + botão ☰ (direita) na mesma linha
_col_titulo, _col_gear = st.columns([11, 1], vertical_alignment="center")
with _col_titulo:
    _titulo_placeholder = st.empty()
with _col_gear:
    if st.button("☰", key="btn_gear"):
        st.session_state.show_admin = not st.session_state.show_admin
        st.rerun()

# ─────────────────────────── PAINEL ADMIN ────────────────────────

if st.session_state.show_admin:
    with st.container(border=True):
        st.markdown("#### ⚙️ Área Admin")

        if st.session_state.admin_logado:
            col_ok, col_sair = st.columns([4, 1])
            col_ok.success("Admin conectado")
            if col_sair.button("Sair", key="admin_sair"):
                st.session_state.admin_logado = False
                st.session_state.show_admin = False
                st.rerun()

            contratos_admin = listar_contratos()
            if not contratos_admin:
                st.info("Nenhum contrato cadastrado ainda.")
                with st.form("form_novo_contrato"):
                    nc_num   = st.text_input("Nº do Contrato")
                    nc_senha = st.text_input("Senha do Contrato", type="password")
                    nc_ide   = st.text_input("Empreendimento", placeholder="ex: SRGE/SI-III/HDTON/CMUGH")
                    if st.form_submit_button("Criar Contrato"):
                        if nc_num.strip() and nc_senha.strip():
                            if criar_contrato(nc_num.strip(), nc_senha.strip(), nc_ide.strip()):
                                st.success("Contrato criado.")
                                st.rerun()
                            else:
                                st.error("Contrato já existe.")
                        else:
                            st.error("Preencha todos os campos.")
            else:
                contrato_admin = st.selectbox("Contrato", contratos_admin,
                                              key="sel_contrato_admin")
                tab_c, tab_f, tab_u, tab_r = st.tabs(["Contratos", "Fiscais", "Unidades", "Registros"])

                # ── Contratos ──
                with tab_c:
                    st.markdown("**Novo Contrato**")
                    with st.form("form_novo_contrato"):
                        nc_num   = st.text_input("Nº do Contrato")
                        nc_senha = st.text_input("Senha do Contrato", type="password")
                        nc_ide   = st.text_input("Empreendimento", placeholder="ex: SRGE/SI-III/HDTON/CMUGH")
                        if st.form_submit_button("Criar"):
                            if nc_num.strip() and nc_senha.strip():
                                if criar_contrato(nc_num.strip(), nc_senha.strip(), nc_ide.strip()):
                                    st.success("Contrato criado.")
                                    st.rerun()
                                else:
                                    st.error("Contrato já existe.")
                            else:
                                st.error("Preencha todos os campos.")

                    st.markdown("**Empreendimento do contrato selecionado**")
                    _ide_atual = obter_identificador(contrato_admin)
                    with st.form("form_ide"):
                        _ide_novo = st.text_input("Empreendimento", value=_ide_atual,
                                                   placeholder="ex: SRGE/SI-III/HDTON/CMUGH")
                        if st.form_submit_button("Salvar"):
                            atualizar_identificador(contrato_admin, _ide_novo.strip())
                            st.success("Empreendimento atualizado.")
                            st.rerun()

                    st.markdown("**Contratos cadastrados**")
                    for _c in contratos_admin:
                        c1, c2 = st.columns([4, 1])
                        c1.write(_c)
                        if c2.button("🗑", key=f"del_c_{_c}"):
                            excluir_contrato(_c)
                            st.rerun()

                # ── Fiscais ──
                with tab_f:
                    st.markdown("**Cadastrar Fiscal**")
                    with st.form("form_fiscal"):
                        f_nome  = st.text_input("Nome do Fiscal de Campo")
                        f_chave = st.text_input("Chave")
                        f_disc  = st.selectbox("Disciplina", DISCIPLINAS, key="disc_fiscal")
                        f_email = st.text_input("E-mail do Fiscal")
                        if st.form_submit_button("Adicionar"):
                            if f_nome.strip():
                                adicionar_fiscal(contrato_admin, f_nome.strip(),
                                                 f_chave.strip(), f_disc, f_email.strip())
                                st.success("Fiscal adicionado.")
                                st.rerun()
                            else:
                                st.error("Informe o nome do fiscal.")
                    st.markdown("**Fiscais cadastrados**")
                    for fiscal in listar_fiscais(contrato_admin):
                        c1, c2 = st.columns([4, 1])
                        c1.write(f"{fiscal['nome']} | {fiscal.get('email') or fiscal['chave']}")
                        if c2.button("🗑", key=f"del_f_{fiscal['id']}"):
                            excluir_fiscal(fiscal["id"])
                            st.rerun()

                # ── Unidades ──
                with tab_u:
                    st.markdown("**Cadastrar Unidade**")
                    with st.form("form_unidade"):
                        u_nome = st.text_input("Nome da Unidade")
                        if st.form_submit_button("Adicionar"):
                            if u_nome.strip():
                                adicionar_unidade(contrato_admin, u_nome.strip())
                                st.success("Unidade adicionada.")
                                st.rerun()
                            else:
                                st.error("Informe o nome da unidade.")
                    st.markdown("**Unidades cadastradas**")
                    for unidade in listar_unidades(contrato_admin):
                        c1, c2 = st.columns([4, 1])
                        c1.write(unidade["nome"])
                        if c2.button("🗑", key=f"del_u_{unidade['id']}"):
                            excluir_unidade(unidade["id"])
                            st.rerun()

                # ── Registros ──
                with tab_r:
                    _regs_admin = carregar_lista(contrato_admin)
                    if not _regs_admin:
                        st.info("Nenhum registro para este contrato.")
                    else:
                        st.markdown(f"**{len(_regs_admin)} registro(s) encontrado(s)**")

                        # seleção para exclusão em lote
                        _sel_ids = []
                        for _r in _regs_admin:
                            _rid_label = id_registro(_r)
                            _ca, _cb = st.columns([1, 9])
                            _checked = _ca.checkbox("", key=f"adm_sel_{_r.id}", label_visibility="collapsed")
                            _cb.markdown(
                                f"**{_rid_label}**  \n"
                                f"<small>{fmt_data(_r.data_registro)} · {_r.disciplina} · {_r.status}</small>",
                                unsafe_allow_html=True,
                            )
                            if _checked:
                                _sel_ids.append(_r.id)

                        st.divider()
                        _col_del, _col_all = st.columns(2)

                        with _col_del:
                            if _sel_ids:
                                if st.button(f"🗑 Excluir selecionados ({len(_sel_ids)})",
                                             type="primary", key="btn_del_sel"):
                                    excluir_registros(_sel_ids)
                                    st.success(f"{len(_sel_ids)} registro(s) excluído(s).")
                                    st.rerun()
                            else:
                                st.caption("Marque registros acima para excluir.")

                        with _col_all:
                            if st.button("🗑 Limpar TODOS os registros", key="btn_del_all"):
                                st.session_state["confirmar_limpar"] = True
                            if st.session_state.get("confirmar_limpar"):
                                st.warning("Tem certeza? Isso apagará **todos** os registros deste contrato.")
                                _cc1, _cc2 = st.columns(2)
                                if _cc1.button("Sim, apagar tudo", type="primary", key="btn_confirmar_limpar"):
                                    excluir_registros([_r.id for _r in _regs_admin])
                                    st.session_state["confirmar_limpar"] = False
                                    st.success("Todos os registros foram apagados.")
                                    st.rerun()
                                if _cc2.button("Cancelar", key="btn_cancelar_limpar"):
                                    st.session_state["confirmar_limpar"] = False
                                    st.rerun()

        else:
            with st.form("admin_login"):
                _user  = st.text_input("Usuário")
                _senha = st.text_input("Senha", type="password")
                if st.form_submit_button("Entrar"):
                    if _verificar_admin(_user.strip(), _senha):
                        st.session_state.admin_logado = True
                        st.rerun()
                    else:
                        st.error("Usuário ou senha inválidos.")

# ─────────────────────────── SELEÇÃO DE CONTRATO ─────────────────

st.subheader("Contrato nº.:")
try:
    contratos_disponiveis = listar_contratos()
except _SupabaseError as _e:
    st.error(str(_e))
    st.stop()

if contratos_disponiveis:
    tenant = st.selectbox("Contrato", contratos_disponiveis, label_visibility="collapsed")
else:
    tenant = st.text_input("Contrato", placeholder="Nº do contrato (solicite ao admin o cadastro)", label_visibility="collapsed")
    if not tenant.strip():
        st.warning("Nenhum contrato cadastrado. Acesse a Área Admin no menu lateral para cadastrar.")
        st.stop()
    tenant = tenant.strip()

try:
    lista_registros = carregar_lista(tenant)
    _empreendimento = obter_identificador(tenant)
except _SupabaseError as _e:
    st.error(str(_e))
    st.stop()
_emp = (_empreendimento or "SRGE/SI-III/HDTON/CMUGH").upper()
_titulo_placeholder.markdown(
    f'<div style="margin:4px 0 8px">'
    f'<p style="font-size:0.85rem;font-weight:400;letter-spacing:0.04em;margin:0;'
    f'color:#555;text-transform:uppercase">{_emp}</p>'
    f'<p style="font-size:1.1rem;font-weight:700;letter-spacing:0.03em;margin:0;'
    f'color:#0D0D0D">RO - REGISTRO DE OCORRÊNCIAS</p>'
    f'</div>',
    unsafe_allow_html=True,
)

# ─────────────────────────── 1) FISCALIZAÇÃO DE CAMPO ───────────

st.subheader("1) Fiscalização de Campo")

fiscais_disponiveis = listar_fiscais(tenant)
fiscal_selecionado  = {}

if fiscais_disponiveis:
    nomes = [""] + [f["nome"] for f in fiscais_disponiveis]
    nome_escolhido = st.selectbox("Nome do Fiscal de Campo", nomes, key="sel_fiscal",
                                  format_func=lambda x: "— selecione —" if x == "" else x)
    fiscal_selecionado = next((f for f in fiscais_disponiveis if f["nome"] == nome_escolhido), {})
    if nome_escolhido:
        col_chave, col_disc_f = st.columns(2)
        col_chave.text_input("Chave",      value=fiscal_selecionado.get("chave", ""),      disabled=True, key="chave_ro")
        col_disc_f.text_input("Disciplina", value=fiscal_selecionado.get("disciplina", ""), disabled=True, key="disc_ro")
else:
    st.info("Nenhum fiscal cadastrado para este contrato. Solicite ao admin.")
    col_nome, col_email_f = st.columns([3, 2])
    with col_nome:
        fiscal_selecionado["nome"] = st.text_input("Nome do Fiscal de Campo", key="fiscal_nome_livre")
    with col_email_f:
        fiscal_selecionado["email"] = st.text_input("E-mail", key="fiscal_email_livre")
    col_chave, col_disc_f = st.columns([2, 3])
    with col_chave:
        fiscal_selecionado["chave"] = st.text_input("Chave", key="fiscal_chave_livre")
    with col_disc_f:
        fiscal_selecionado["disciplina"] = st.selectbox("Disciplina", DISCIPLINAS, key="fiscal_disc_livre")

fiscal_nome  = fiscal_selecionado.get("nome", "")
fiscal_chave = fiscal_selecionado.get("chave", "")

# ─────────────────────────── 2) REGISTROS RDOe ───────────────────

st.subheader("2) Registros de Ocorrências")

if "fk" not in st.session_state:
    st.session_state.fk = 0
fk = st.session_state.fk

# Unidade
_SELECIONE = "— selecione —"
unidades_disponiveis = [u["nome"] for u in listar_unidades(tenant)]
col_unidade, col_data = st.columns([4, 1])
with col_unidade:
    if unidades_disponiveis:
        obra = st.selectbox("Unidade", [_SELECIONE] + unidades_disponiveis, key=f"obra_{fk}")
    else:
        obra = st.text_input("Unidade", key=f"obra_{fk}")
with col_data:
    data_registro = st.date_input("Data", value=None, format="DD/MM/YYYY", key=f"data_{fk}")

col_frente, col_disciplina = st.columns([3, 2])
with col_frente:
    frente_servico = st.text_input("Frente de serviço", key=f"frente_{fk}")
with col_disciplina:
    disciplina = st.selectbox("Disciplina", [_SELECIONE] + DISCIPLINAS, key=f"disc_{fk}")

atividade = st.text_area("Atividade Executada", key=f"ativ_{fk}")

fotos = st.file_uploader(
    "Evidências (máx. 4 fotos, 20 MB cada)",
    type=["jpg", "jpeg", "png", "webp"],
    accept_multiple_files=True,
    key=f"fotos_{fk}",
)
fotos_validas = []
for f in fotos:
    if len(f.getvalue()) > 20 * 1024 * 1024:
        st.warning(f"'{f.name}' excede 20 MB e foi ignorada.")
    else:
        fotos_validas.append(f)
fotos = fotos_validas[:4]
if len(fotos_validas) > 4:
    st.warning("Apenas as 4 primeiras fotos serão consideradas.")

_legendas: list[str] = []
if fotos:
    _thumbs = [img_thumb(f.getvalue()) for f in fotos]
    if len(_thumbs) == 1:
        _img_centralizada(_thumbs[0])
        _legendas.append(st.text_input("Legenda", key=f"cap_0_{fk}", placeholder="Descrição da evidência", label_visibility="collapsed"))
    else:
        for _i in range(0, len(_thumbs), 2):
            _cols = st.columns(2)
            for _j in range(2):
                _idx = _i + _j
                if _idx < len(_thumbs):
                    with _cols[_j]:
                        _img_centralizada(_thumbs[_idx])
                        _legendas.append(st.text_input("Legenda", key=f"cap_{_idx}_{fk}", placeholder="Descrição da evidência", label_visibility="collapsed"))

col_equipe, col_resp = st.columns(2)
with col_equipe:
    equipe      = st.text_input("Equipe da Contratada", key=f"equipe_{fk}")
with col_resp:
    responsavel = st.text_input("Responsável Contratada", key=f"resp_{fk}")
status      = st.selectbox("Status", [_SELECIONE, "Executado", "Em andamento", "Bloqueado", "Não iniciado"], key=f"status_{fk}")
impacto_rdo = st.selectbox("Classificação do Registro", [_SELECIONE, "Alta", "Média", "Baixa"], key=f"impacto_{fk}")
observacoes = st.text_area("Observações adicionais", key=f"obs_{fk}")

if st.button("Salvar", type="primary"):
    _obra_val = str(obra).strip()
    _campos_invalidos = []
    if not fiscal_nome.strip():            _campos_invalidos.append("Fiscal de Campo")
    if not _obra_val or _obra_val == _SELECIONE: _campos_invalidos.append("Unidade")
    if data_registro is None:              _campos_invalidos.append("Data")
    if disciplina == _SELECIONE:           _campos_invalidos.append("Disciplina")
    if not frente_servico.strip():         _campos_invalidos.append("Frente de serviço")
    if not atividade.strip():              _campos_invalidos.append("Atividade Executada")
    if status == _SELECIONE:              _campos_invalidos.append("Status")
    if impacto_rdo == _SELECIONE:         _campos_invalidos.append("Classificação do Registro")
    if _campos_invalidos:
        st.error(f"Preencha os campos obrigatórios: {', '.join(_campos_invalidos)}")
    else:
        salvar_registro(RegistroCM(
            id=0,
            tenant=tenant,
            data_registro=data_registro.isoformat(),
            obra=str(obra).strip(),
            frente_servico=frente_servico.strip(),
            disciplina=disciplina,
            atividade=atividade.strip(),
            equipe=equipe.strip(),
            responsavel=responsavel.strip(),
            fiscal=fiscal_nome.strip(),
            status=status,
            impacto_rdo=impacto_rdo,
            observacoes=observacoes.strip(),
            evidencias=json.dumps([
                {"foto": base64.b64encode(f.getvalue()).decode(), "legenda": leg}
                for f, leg in zip(fotos[:4], _legendas)
            ]),
            chave=fiscal_chave.strip(),
        ))
        st.success("Registro inserido com sucesso.")
        st.session_state.fk += 1
        st.rerun()

# ─────────────────────────── 3) LIST DO FISCAL ───────────────────

st.subheader("3) Registros")

if lista_registros:
    exibir_cards(lista_registros, tenant, _empreendimento)
else:
    st.info("Nenhum registro encontrado para este contrato.")

if lista_registros:
    selecionados = [r for r in lista_registros if st.session_state.get(f"sel_{r.id}", False)]
    if selecionados:
        _n_sel = f"{len(selecionados)}-selecionados"
        _lbl = f"PDF dos selecionados ({len(selecionados)})"
        _share_btn(_lbl, _to_pdf(selecionados, tenant, _empreendimento),
                   f"RO-{_n_sel}.pdf", "application/pdf", f"pdf_sel_{_n_sel}")

    _n = f"{len(lista_registros):04d}-registros"
    _share_btn("Compartilhar Excel (todos os registros)", _to_excel(lista_registros),
               f"RO-{_n}.xlsx",
               "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
               f"xlsx_{_n}")
